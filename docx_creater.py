import docx
doc=docx.Document()
doc.add_paragraph('Hello World!')
paraobj1=doc.add_paragraph('This is second paragraph.')
paraobj2=doc.add_paragraph('This is third paragraph.')
paraobj1.add_run('This text i am adding at the end of the first paragrapg.')
doc.save('multipleparagraphs.docx')
