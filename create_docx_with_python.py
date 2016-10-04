import docx
#create document object
doc=docx.Document()
#adding paragraph to document
doc.add_paragraph('Hello World!')
paraobj1=doc.add_paragraph('This is second paragraph.')
paraobj2=doc.add_paragraph('This is third paragraph.')
#adding this line to end of first paragraph. 
paraobj1.add_run('This text i am adding at the end of the first paragraph.')
#saving the document
doc.save('multipleparagraphs.docx')
